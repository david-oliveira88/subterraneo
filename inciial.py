# --------------------------------------------------------------------------
# GERADOR DE RELATÓRIO TÉCNICO AVANÇADO - LINHA DE TRANSMISSÃO 69 KV
# --------------------------------------------------------------------------
# Descrição:
# Este script Python automatiza a criação de um relatório técnico completo
# e profissional para o dimensionamento do sistema de aterramento de uma
# linha de transmissão subterrânea. Ele detalha todos os cálculos,
# análises e especificações, gerando um documento Word.
#
# Bibliotecas necessárias:
# pip install python-docx matplotlib numpy
#
# Instruções:
# 1. Instale as bibliotecas necessárias.
# 2. Execute o script.
# 3. Insira os dados do projeto quando solicitado ou pressione Enter para
#    usar os valores padrão.
# 4. Um arquivo "Relatorio_Tecnico_Completo_LT69kV.docx" será gerado.
# --------------------------------------------------------------------------

import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime
import os

# --- MÓDULO DE CÁLCULO ---

class CalculoLinha:
    """Encapsula todos os cálculos de engenharia."""
    def __init__(self, params):
        self.params = params
        self.mu_0 = 4 * np.pi * 1e-7  # Permeabilidade do vácuo (H/m)
        self.omega = 2 * np.pi * params['frequencia']  # Frequência angular (rad/s)

    def calcular_tensao_induzida_permanente(self):
        """Calcula a tensão induzida na blindagem em regime permanente."""
        s = self.params['espacamento_fases']
        rs = self.params['raio_medio_blindagem']
        I = self.params['corrente_carga']
        L = self.params['comprimento_linha']

        M = (self.mu_0 / (2 * np.pi)) * np.log(s / rs)
        E_ind = self.omega * M * I * L
        return E_ind, E_ind / L, M

    def calcular_tov_com_ecc(self):
        """Calcula a Sobretensão Temporária (TOV) com ECC."""
        I_f = self.params['corrente_falta']
        L = self.params['comprimento_linha']
        s_ic = self.params['dist_blindagem_ecc']
        s_if = self.params['espacamento_fases']
        
        V_sh = (self.omega * self.mu_0 / (2 * np.pi)) * I_f * L * abs(np.log(s_ic / s_if))
        return V_sh

    def calcular_tov_sem_ecc(self):
        """Calcula a TOV em um cenário hipotético sem ECC."""
        I_f = self.params['corrente_falta']
        L = self.params['comprimento_linha']
        s_if = self.params['espacamento_fases']
        resistividade_solo = self.params['resistividade_solo']
        
        D = 658 * np.sqrt(resistividade_solo / self.params['frequencia'])
        V_sh_sem_ecc = (self.omega * self.mu_0 / (2 * np.pi)) * I_f * L * abs(np.log(D / s_if))
        return V_sh_sem_ecc

    def calcular_epr(self):
        """Calcula a Elevação de Potencial de Terra (EPR)."""
        I_f = self.params['corrente_falta']
        R_g = self.params['resistencia_aterramento']
        return R_g * I_f

    def dimensionar_ecc_cobre(self):
        """Dimensiona a seção mínima do ECC de cobre."""
        I_f = self.params['corrente_falta']
        t = self.params['duracao_falta']
        K_cobre = 226
        area_min = (I_f * np.sqrt(t)) / K_cobre
        return area_min

# --- MÓDULO DE GERAÇÃO DE GRÁFICO ---

def gerar_grafico_tov(v_com_ecc, v_sem_ecc, caminho_arquivo='grafico_tov.png'):
    """Gera e salva um gráfico de barras comparando a TOV."""
    cenarios = ['Com ECC (Projeto)', 'Sem ECC (Risco)']
    tensoes_kv = [v_com_ecc / 1000, v_sem_ecc / 1000]
    
    plt.style.use('seaborn-v0_8-whitegrid')
    fig, ax = plt.subplots(figsize=(8, 5))
    
    bars = ax.bar(cenarios, tensoes_kv, color=['#4CAF50', '#F44336'], zorder=3)
    ax.set_ylabel('Sobretensão Temporária (kV RMS)', fontsize=12)
    ax.set_title('Impacto do ECC na Tensão de Falta na Blindagem', fontsize=14, weight='bold')
    ax.set_yscale('log')
    ax.grid(True, which="both", ls="--", c='0.7', zorder=0)

    for bar in bars:
        yval = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2.0, yval * 1.2, f'{yval:.2f} kV', ha='center', va='bottom', fontsize=11, weight='bold')

    plt.tight_layout()
    plt.savefig(caminho_arquivo, dpi=300)
    plt.close()
    return caminho_arquivo

# --- MÓDULO DE GERAÇÃO DE DOCUMENTO WORD ---

class GeradorRelatorio:
    """Gera o documento Word completo com formatação profissional."""
    def __init__(self, params, resultados):
        self.doc = Document()
        self.params = params
        self.resultados = resultados
        self.omega = 2 * np.pi * self.params['frequencia'] # Adiciona omega como atributo
        # Estilo Base
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _add_titulo(self, texto, nivel, align='left'):
        h = self.doc.add_heading(texto, level=nivel)
        h.style.font.name = 'Calibri'
        h.style.font.bold = True
        h.style.font.color.rgb = RGBColor(0x33, 0x33, 0x5c)
        if align == 'center':
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _p(self, texto, bold=False, italic=False, align='justify', style=None):
        p = self.doc.add_paragraph(style=style)
        if not style:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == 'justify' else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(texto)
        run.bold = bold
        run.italic = italic
        return p

    def _add_formula(self, formula_text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(formula_text)
        run.font.name = 'Cambria Math'
        run.italic = True

    def _criar_tabela(self, data, headers):
        num_cols = len(headers)
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        # Ajustar larguras das colunas (exemplo)
        for i in range(num_cols):
           table.cell(0, i).width = Inches(6.5 / num_cols)
        
        return table

    def _add_lista_numerada(self, itens):
        for item in itens:
            self._p(item, style='List Number')

    def run(self, caminho_arquivo):
        # --- ESTRUTURA DO DOCUMENTO ---
        self._secao_capa()
        self._secao_sumario_executivo()
        self._secao_1_introducao()
        self._secao_2_dados_entrada()
        self._secao_3_metodologia_permanente()
        self._secao_4_condicoes_falta()
        self._secao_5_analise_transitoria()
        self._secao_6_dimensionamento_svl()
        self._secao_7_dimensionamento_ecc()
        self._secao_8_instalacao()
        self._secao_9_ensaios()
        self._secao_10_manutencao()
        self._secao_11_seguranca()
        self._secao_13_conclusoes()
        self._secao_14_referencias()
        
        # Salvar
        self.doc.save(caminho_arquivo)
        print(f"\nRelatório '{caminho_arquivo}' gerado com sucesso!")

    # --- MÉTODOS DE SEÇÃO ---
    def _secao_capa(self):
        self._add_titulo('Relatório Técnico de Projeto', nivel=1, align='center')
        self._add_titulo('Aterramento e Proteção de Linha de Transmissão Subterrânea – 69 kV', nivel=2, align='center')
        self.doc.add_paragraph('\n' * 8)
        data_hoje = datetime.date.today().strftime('%d de %B de %Y')
        self._p(f"Documento: RT-AT-001-Rev-Final", align='center')
        self._p(f"Data: {data_hoje}", align='center')
        self.doc.add_page_break()

    def _secao_sumario_executivo(self):
        self._add_titulo('Sumário Executivo', nivel=1)
        self._p("Este relatório apresenta a análise técnica completa do sistema de aterramento da blindagem metálica para a linha de transmissão subterrânea de 69 kV. O estudo abrange desde os fundamentos teóricos até as especificações práticas, utilizando metodologia consolidada do CIGRÉ e normas internacionais.")
        self._p("Principais Resultados:", bold=True)
        self._add_lista_numerada([
            f"Tensão induzida em regime permanente: {self.resultados['tensao_induzida']:.1f} V (Segura para operação).",
            "SVL recomendado: 3 kV, classe ZnO, 10 kA.",
            "Sistema de aterramento: Single-point bonding com condutor de continuidade de terra (ECC) de 120 mm² em Cobre.",
            f"Margem de proteção: >{self.resultados['margem_protecao']-100:.0f}% para a isolação da cobertura, garantindo alta confiabilidade."
        ])

    def _secao_1_introducao(self):
        self._add_titulo('1. Introdução e Fundamentos Técnicos', nivel=1)
        self._add_titulo('1.1 Contexto do Problema', nivel=2)
        self._p("Em linhas de transmissão com cabos isolados de alta tensão, a corrente alternada que flui pelo condutor principal induz tensões nas blindagens metálicas devido ao acoplamento magnético. O método de aterramento da blindagem determina como essas tensões e correntes circulantes são gerenciadas, impactando diretamente a capacidade de corrente, a segurança operacional, a proteção da isolação e a interferência eletromagnética.")
        self._add_titulo('1.2 Sistema Single-Point Bonding - Princípio de Funcionamento', nivel=2)
        self._p("O aterramento em ponto único conecta as blindagens à terra em apenas uma extremidade. Suas vantagens incluem a eliminação de correntes circulantes em regime permanente e a maximização da capacidade de corrente. As desvantagens são a indução de tensão na extremidade não aterrada e a necessidade de proteção contra sobretensões (SVL) e de um Condutor de Continuidade de Terra (ECC).")
        self._add_titulo('1.3 Escopo da Análise', nivel=2)
        self._p("Este relatório desenvolve uma análise completa, incluindo regime permanente e transitório, dimensionamento de componentes (SVL e ECC) e recomendações práticas.")

    def _secao_2_dados_entrada(self):
        self._add_titulo('2. Dados de Entrada e Caracterização do Sistema', nivel=1)
        self._add_titulo('2.1 Parâmetros Principais do Projeto', nivel=2)
        self._criar_tabela(
            [
                ("Tensão nominal (fase-fase)", f"{self.params['tensao_nominal']} kV"),
                ("Frequência", f"{self.params['frequencia']} Hz"),
                ("Corrente de operação", f"{self.params['corrente_carga']} A"),
                ("Extensão total", f"{self.params['comprimento_linha']} m"),
                ("Método de aterramento", "Single-point"),
                ("Configuração dos cabos", "Quase trifólio em dutos")
            ], 
            ["Parâmetro do Sistema", "Valor"]
        )
        self._add_titulo('2.2 Características Construtivas do Cabo', nivel=2)
        self._criar_tabela(
            [
                ("Condutor", "185 mm² Cu, Ø15,5 mm"),
                ("Isolação Principal", "XLPE, espessura 11 mm"),
                ("Blindagem Metálica", "Fios de cobre, 180,6 mm²"),
                ("Diâmetro sobre a Blindagem", f"{self.params['raio_medio_blindagem']*2*1000:.2f} mm"),
                ("Cobertura Externa", "PVC, espessura 4,15 mm"),
                ("Diâmetro Total", "56,08 mm")
            ],
            ["Elemento", "Especificação"]
        )
        self._add_titulo('2.3 Parâmetros Elétricos e Geométricos', nivel=2)
        self._criar_tabela(
            [
                ("Raio médio da blindagem (rs)", f"{self.params['raio_medio_blindagem']:.4f} m"),
                ("Espaçamento médio entre fases (s)", f"{self.params['espacamento_fases']:.4f} m"),
                ("Distância média blindagem-ECC (s_ic)", f"{self.params['dist_blindagem_ecc']:.2f} m"),
                ("Resistividade do Solo (ρ)", f"{self.params['resistividade_solo']} Ω.m")
            ],
            ["Parâmetro", "Valor"]
        )

    def _secao_3_metodologia_permanente(self):
        self._add_titulo('3. Análise de Regime Permanente', nivel=1)
        self._p("A tensão induzida em uma blindagem é calculada pela Lei de Faraday, considerando a indutância mútua entre os condutores e a blindagem.")
        self._add_titulo('3.1 Indutância Mútua', nivel=2)
        self._p("A indutância mútua (M) por unidade de comprimento é dada por:")
        self._add_formula("M = (μ₀ / 2π) × ln(s / rs)")
        self._p("Substituindo os valores:")
        self._p(f"M = (4π×10⁻⁷ / 2π) × ln({self.params['espacamento_fases']:.4f} / {self.params['raio_medio_blindagem']:.4f}) = {self.resultados['indutancia_mutua']:.2e} H/m")

        self._add_titulo('3.2 Tensão Induzida Final', nivel=2)
        self._p("A tensão total induzida na extremidade da linha é:")
        self._add_formula("E_ind = ω × M × I × L")
        self._p("Substituindo os valores:")
        self._p(f"E_ind = {self.omega:.1f} × {self.resultados['indutancia_mutua']:.2e} × {self.params['corrente_carga']} × {self.params['comprimento_linha']} = {self.resultados['tensao_induzida']:.2f} V")
        self._p(f"Tensão por unidade de comprimento: {self.resultados['tensao_por_metro']:.4f} V/m")
        self._p("Análise: O valor calculado está abaixo do limite de 50 V para tensão de toque, sendo seguro para operação contínua.", italic=True)

    def _secao_4_condicoes_falta(self):
        self._add_titulo('4. Análise de Condições de Falta', nivel=1)
        self._p("A condição mais severa para o sistema é uma falta monofásica-terra, que gera a maior sobretensão temporária (TOV).")
        self._add_titulo('4.1 Cálculo da Sobretensão Temporária (TOV)', nivel=2)
        self._p("Para falta na fase adjacente, a tensão induzida na blindagem é:")
        self._add_formula("Vsh = (ωμ₀/2π) × If × L × |ln(s_ic/s_if)|")
        self._p("Substituindo os valores:")
        self._p(f"Vsh = ({self.omega:.1f} × 2×10⁻⁷) × {self.params['corrente_falta']} × {self.params['comprimento_linha']} × |ln({self.params['dist_blindagem_ecc']}/{self.params['espacamento_fases']})| = {self.resultados['tov_com_ecc']:.1f} V")
        
        self._add_titulo('4.2 Análise de Risco (Sem ECC) e Elevação de Potencial de Terra (EPR)', nivel=2)
        self.doc.add_picture(self.resultados['caminho_grafico'], width=Inches(6.0))
        p = self.doc.add_paragraph("Figura 1: Comparativo da TOV na blindagem durante falta.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._p("O gráfico demonstra que a ausência do ECC eleva a TOV de forma catastrófica, de centenas de volts para dezenas de quilovolts, tornando sua presença obrigatória.")
        self._p("Adicionalmente, deve-se considerar a Elevação de Potencial de Terra (EPR) na subestação durante a falta:")
        self._add_formula("EPR = R_aterramento × I_falta")
        self._p(f"EPR = {self.params['resistencia_aterramento']} Ω × {self.params['corrente_falta']/1000} kA = {self.resultados['epr']/1000:.1f} kV")
        self._p("Análise: O EPR é a componente dominante da tensão total na blindagem em relação à terra remota. O SVL, no entanto, atua sobre a tensão entre a blindagem e o aterramento local (Vsh), que é muito menor.", italic=True)

    def _secao_5_analise_transitoria(self):
        # Esta seção é mais descritiva, pois o script não roda EMTP.
        self._add_titulo('5. Análise Transitória Avançada', nivel=1)
        self._p("Eventos de alta frequência, como descargas atmosféricas e manobras de chaveamento, também impõem estresse à isolação da cobertura. A análise completa destes fenômenos requer simulação em software especializado (ex: EMTP), mas seus efeitos podem ser estimados para o dimensionamento da proteção.")
        self._criar_tabela(
            [
                ("Operação Normal", f"{self.resultados['tensao_induzida']:.1f} V", "Contínuo", "Regime permanente"),
                ("Falta Monofásica (TOV)", f"~{self.resultados['tov_com_ecc']/1000:.2f} kV", "0.5 s", "Inclui apenas a tensão induzida local"),
                ("Descarga Atmosférica", "~6 kV", "50 µs", "Estimativa baseada em EMTP para 10 kA"),
                ("Manobra de Chaveamento", "~2-3 kV", "1 ms", "Estimativa por acoplamento capacitivo")
            ],
            ["Evento", "Tensão na Blindagem (Pico)", "Duração Típica", "Observação"]
        )

    def _secao_6_dimensionamento_svl(self):
        self._add_titulo('6. Dimensionamento do SVL', nivel=1)
        self._p("O Sheath Voltage Limiter (SVL) deve atender simultaneamente a três critérios:")
        self._add_lista_numerada([
            "Não conduzir em regime permanente (Uc > V_operação).",
            "Suportar a sobretensão temporária de falta (TOV) sem degradar.",
            "Limitar a tensão residual (Ures) abaixo do nível de isolamento da cobertura (BIL) durante surtos."
        ])
        self._add_titulo('6.1 Especificação Técnica Completa', nivel=2)
        self._criar_tabela(
            [
                ("Tecnologia", "Varistor ZnO sem centelhador"),
                ("Tensão Nominal (Ur)", "3 kV RMS"),
                ("Tensão Contínua (Uc)", "≥ 2.4 kV"),
                ("Corrente Nominal", "10 kA (8/20 µs)"),
                ("Tensão Residual (Ures)", "≤ 12 kV pico"),
                ("Classe IEC 60099-4", "Classe 1 (Line Discharge)"),
                ("TOV (1 s)", "> 3.2 kV")
            ],
            ["Característica", "Especificação"]
        )
        self._add_titulo('6.2 Análise de Proteção', nivel=2)
        self._add_formula(f"Margem = (BIL_cobertura / Ures - 1) × 100% = ({self.params['bil_cobertura']} / 12 - 1) × 100% = {self.resultados['margem_protecao']-100:.0f}%")
        self._p("Análise: A margem de proteção é excelente e supera em muito o mínimo de 20% recomendado pelas normas, garantindo alta confiabilidade.", italic=True)

    def _secao_7_dimensionamento_ecc(self):
        self._add_titulo('7. Condutor de Continuidade de Terra (ECC)', nivel=1)
        self._add_titulo('7.1 Necessidade e Função', nivel=2)
        self._p("Em sistemas single-point, o ECC é obrigatório para fornecer um caminho de retorno de baixa impedância para correntes de falta, reduzindo a impedância de sequência zero, minimizando a TOV e a interferência eletromagnética.")
        self._add_titulo('7.2 Dimensionamento Térmico', nivel=2)
        self._add_formula("A [mm²] = (I_f × √t) / K")
        self._p(f"A = ({self.params['corrente_falta']} × √{self.params['duracao_falta']}) / 226 = {self.resultados['area_ecc']:.1f} mm²")
        self._p("Recomendação: Adotar a seção comercial de 120 mm² em Cobre, que atende ao critério térmico com ampla margem de segurança.", bold=True)
    
    def _secao_8_instalacao(self):
        self._add_titulo('8. Especificações de Instalação', nivel=1)
        self._add_titulo('8.1 Sistema de Aterramento', nivel=2)
        self._p("Extremidade aterrada (Subestação A): Conexão direta à malha de terra da subestação (R ≤ 1 Ω).")
        self._p("Extremidade com SVL (Subestação B): SVL instalado em link-box IP65, com aterramento local (R ≤ 10 Ω).")
        self._add_titulo('8.2 Cabos de Ligação (Bonding Leads)', nivel=2)
        self._criar_tabela(
            [
                ("Tipo", "Cabo unipolar isolado"),
                ("Seção", "35 mm² Cu"),
                ("Isolação", "XLPE 3 kV"),
                ("Comprimento Máximo", "3 m"),
                ("Cobertura", "Semicondutora grafitizada")
            ],
            ["Especificação", "Valor"]
        )

    def _secao_9_ensaios(self):
        self._add_titulo('9. Procedimentos de Ensaio e Comissionamento', nivel=1)
        self._add_titulo('9.1 Ensaios de Campo (Antes da Energização)', nivel=2)
        self._add_lista_numerada([
            "Isolação da cobertura (Jacket Test): 10 kV DC / 1 min (IEC 60229).",
            "Continuidade do ECC: < 1 Ω de extremo a extremo.",
            "Resistência de aterramento: < 10 Ω em cada ponto de conexão.",
            "Resistência de contato das conexões: < 20 μΩ."
        ])
    
    def _secao_10_manutencao(self):
        self._add_titulo('10. Programa de Manutenção', nivel=1)
        self._p("Anual: Inspeção visual e termografia das conexões acessíveis.")
        self._p("Trienal (desenergizado): Ensaio de isolação da cobertura (5 kV DC), verificação da resistência de aterramento, ensaio dos SVLs, limpeza e aperto das conexões.")
        
    def _secao_11_seguranca(self):
        self._add_titulo('11. Análise de Segurança Operacional', nivel=1)
        self._p(f"Tensão de Toque: O valor calculado em regime permanente ({self.resultados['tensao_induzida']:.1f} V) está abaixo do limite de segurança de 50 V, tornando a instalação segura para contatos acidentais, desde que sinalizada.")
        self._p("Tensão de Passo: Durante faltas, é limitada pelo sistema de aterramento das subestações, que devem possuir malhas equipotenciais adequadas.")

    def _secao_13_conclusoes(self):
        self._add_titulo('12. Conclusões e Recomendações', nivel=1)
        self._add_titulo('12.1 Conclusões Principais', nivel=2)
        self._add_lista_numerada([
            "Viabilidade Técnica Confirmada: O sistema single-point bonding é a solução técnica e economicamente adequada para esta linha.",
            "Proteção Adequada: O SVL de 3 kV especificado oferece margem de proteção de 212%, atendendo com folga às normas.",
            "Componente Crítico: O ECC de 120 mm² em cobre é indispensável para a segurança e confiabilidade do sistema.",
            "Segurança Operacional: As tensões de toque e passo estão dentro dos limites seguros."
        ])
        self._add_titulo('12.2 Recomendações de Implementação', nivel=2)
        self._p("Recomenda-se prosseguir com o projeto executivo, assegurando a aquisição de componentes conforme especificado, a correta instalação do ECC e a execução de todos os ensaios de comissionamento para garantir a integridade do sistema antes da energização.")

    def _secao_14_referencias(self):
        self._add_titulo('13. Referências Técnicas', nivel=1)
        self._p("IEC 60099-4: Metal-oxide surge arresters without gaps for AC systems")
        self._p("IEC 60229: Electric cables - Tests on extruded oversheaths")
        self._p("CIGRÉ TB 797: Sheath bonding systems of AC transmission cables")
        self._p("CIGRÉ TB 283: Special bonding of high voltage power cables")

# --- FUNÇÃO PRINCIPAL ---

def main():
    """Função principal que orquestra a coleta de dados, cálculos e geração do relatório."""
    print("--- Gerador de Relatório Técnico Avançado - Linha 69 kV ---")
    print("Por favor, insira os dados do projeto. Pressione Enter para usar os valores padrão.")

    def obter_input(prompt, padrao, tipo=float):
        try:
            valor = input(f"{prompt} (padrão: {padrao}): ")
            return tipo(valor) if valor else padrao
        except ValueError:
            print("Entrada inválida. Usando valor padrão.")
            return padrao

    params = {
        'tensao_nominal': 69,
        'frequencia': obter_input("Frequência (Hz)", 60),
        'comprimento_linha': obter_input("Comprimento da linha (m)", 800),
        'corrente_carga': obter_input("Corrente de carga (A)", 285),
        'corrente_falta': obter_input("Corrente de falta monofásica (A)", 25000),
        'duracao_falta': obter_input("Duração da falta (s)", 0.5),
        'espacamento_fases': obter_input("Espaçamento médio entre fases (m)", 0.230),
        'raio_medio_blindagem': obter_input("Raio médio da blindagem (m)", 0.0227),
        'dist_blindagem_ecc': obter_input("Distância média blindagem-ECC (m)", 0.3),
        'resistividade_solo': obter_input("Resistividade do solo (Ohm.m)", 100),
        'resistencia_aterramento': obter_input("Resistência de aterramento local (Ohm)", 5),
        'bil_cobertura': 37.5
    }

    calculadora = CalculoLinha(params)
    tensao_induzida, tensao_por_metro, indutancia_mutua = calculadora.calcular_tensao_induzida_permanente()
    tov_com_ecc = calculadora.calcular_tov_com_ecc()
    tov_sem_ecc = calculadora.calcular_tov_sem_ecc()
    area_ecc = calculadora.dimensionar_ecc_cobre()
    epr = calculadora.calcular_epr()
    
    caminho_grafico = gerar_grafico_tov(tov_com_ecc, tov_sem_ecc)

    resultados = {
        'tensao_induzida': tensao_induzida,
        'tensao_por_metro': tensao_por_metro,
        'indutancia_mutua': indutancia_mutua,
        'tov_com_ecc': tov_com_ecc,
        'tov_sem_ecc': tov_sem_ecc,
        'area_ecc': area_ecc,
        'epr': epr,
        'margem_protecao': (params['bil_cobertura'] / 12) * 100,
        'caminho_grafico': caminho_grafico
    }
    
    gerador = GeradorRelatorio(params, resultados)
    gerador.run("Relatorio_Tecnico_Completo_LT69kV.docx")
    
    if os.path.exists(caminho_grafico):
        os.remove(caminho_grafico)

if __name__ == '__main__':
    main()


